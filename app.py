import io
import re
from typing import List, Tuple

import streamlit as st
from docx import Document
from docx.shared import Pt

import fitz  # PyMuPDF
from openai import OpenAI


def normalize_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


def extract_pdf_pages_text(pdf_bytes: bytes) -> List[str]:
    pages = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        for i in range(doc.page_count):
            page = doc.load_page(i)
            txt = page.get_text("text") or ""
            pages.append(normalize_whitespace(txt))
    return pages


def build_docx(items: List[Tuple[str, List[Tuple[str, str]]]]) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for file_label, pages in items:
        doc.add_heading(file_label, level=1)
        for page_label, translated in pages:
            doc.add_heading(page_label, level=2)
            for para in translated.split("\n\n"):
                p = para.strip()
                if p:
                    doc.add_paragraph(p)
        doc.add_page_break()

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def translate_text(client: OpenAI, hebrew_text: str, target_style: str, extra: str) -> str:
    if not hebrew_text.strip():
        return ""

    prompt = f"""You are a careful Hebrew-to-English translator.
Translate the text accurately. Do not add commentary.
Preserve paragraph structure.

Target style: {target_style}
Extra instructions: {extra or 'None'}

HEBREW TEXT:
{hebrew_text}
""".strip()

    resp = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt,
    )
    return (resp.output_text or "").strip()


st.set_page_config(page_title="Hebrew Pages → Translation → Word Doc", layout="wide")
st.title("Hebrew Pages → Translation → Export to Word")

st.write("Upload PDFs (best if selectable text). The app translates page-by-page and exports one .docx.")

api_key = st.secrets.get("OPENAI_API_KEY", "").strip()
if not api_key:
    st.error("Missing OPENAI_API_KEY. In Streamlit Cloud: App settings → Secrets → add OPENAI_API_KEY.")
    st.stop()

client = OpenAI(api_key=api_key)

target_style = st.selectbox(
    "English translation style",
    [
        "Clear and straightforward",
        "More literal / closer to Hebrew",
        "Warm / chassidus-tone (still accurate)",
        "Academic / formal",
    ],
    index=0,
)

extra_instructions = st.text_area(
    "Optional instructions",
    value="",
    height=90,
)

uploads = st.file_uploader(
    "Upload PDF files",
    type=["pdf"],
    accept_multiple_files=True,
)

run_btn = st.button("Translate and build Word document", type="primary")

if run_btn:
    if not uploads:
        st.warning("Upload at least one PDF.")
        st.stop()

    results: List[Tuple[str, List[Tuple[str, str]]]] = []

    # Count pages for progress
    total_pages = 0
    pdf_cache = []
    for up in uploads:
        pdf_bytes = up.getvalue()
        pages = extract_pdf_pages_text(pdf_bytes)
        pdf_cache.append((up.name, pages))
        total_pages += len(pages)

    prog = st.progress(0)
    done = 0

    for filename, pages in pdf_cache:
        page_outputs: List[Tuple[str, str]] = []
        for i, heb in enumerate(pages, start=1):
            page_label = f"Page {i}"

            if len(heb.strip()) < 10:
                translated = (
                    "[No selectable text detected on this page.]\n\n"
                    "If this PDF page is a scanned image, you’ll need Hebrew OCR added to the app."
                )
            else:
                translated = translate_text(client, heb, target_style, extra_instructions)

            page_outputs.append((page_label, translated))

            done += 1
            prog.progress(min(done / max(total_pages, 1), 1.0))

        results.append((filename, page_outputs))

    docx_bytes = build_docx(results)

    st.success("Done.")
    st.download_button(
        "Download translation.docx",
        data=docx_bytes,
        file_name="translation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
